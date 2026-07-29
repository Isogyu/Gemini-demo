"""アップロードされたファイルから立論の本文を取り出す。"""

import io
import zipfile
from typing import List

from docx import Document

from app.schemas.debate import ExtractedDocument
from app.services.debate.parser import detect_side, extract_title

SUPPORTED_SUFFIXES = (".docx", ".txt", ".md")


class UnsupportedFileError(ValueError):
    pass


def _docx_to_text(content: bytes) -> str:
    try:
        document = Document(io.BytesIO(content))
    except (zipfile.BadZipFile, KeyError) as exc:
        raise UnsupportedFileError("docx ファイルとして読み取れませんでした") from exc
    lines: List[str] = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            lines.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(line for line in lines if line.strip())


def extract_document(filename: str, content: bytes) -> ExtractedDocument:
    lowered = filename.lower()
    if lowered.endswith(".docx"):
        text = _docx_to_text(content)
    elif lowered.endswith((".txt", ".md")):
        text = content.decode("utf-8", errors="replace")
    else:
        raise UnsupportedFileError(
            f"対応していない拡張子です（対応: {', '.join(SUPPORTED_SUFFIXES)}）"
        )
    if not text.strip():
        raise UnsupportedFileError("本文を抽出できませんでした")
    return ExtractedDocument(
        title=extract_title(text),
        text=text,
        detected_side=detect_side(text),
    )
